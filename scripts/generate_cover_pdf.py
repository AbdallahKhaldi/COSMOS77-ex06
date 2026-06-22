"""Fill the lecturer's cover template and convert it to PDF (Phase 13).

Targets HW6 (exercise number = 6, ex06 repo URL). Field values come from the
pure, unit-tested :func:`build_field_values`; :func:`fill_docx` appends a run to
each label's paragraph (the lecturer forbids moving fields). Run with:
``uvx --with python-docx --with docx2pdf python scripts/generate_cover_pdf.py
--template <docx> --output ~/COSMOS77/HW6/COSMOS77-ex06.pdf --exercise-number 6``.
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from collections.abc import Iterable
from pathlib import Path

_REPO_URL = "https://github.com/AbdallahKhaldi/COSMOS77-ex06"
_STUDENTS = (
    ("Student 1", ("212389712", "Abdallah", "Khaldi", "עבדאללה", "חאלדי")),
    ("Student 2", ("323118794", "Tasneem", "Natour", "תסנים", "נאטור")),
)
_STUDENT_LABELS = (
    "ID card",
    "First name in English",
    "Last name in English",
    "First name in Hebrew",
    "Last name in Hebrew",
)


def build_field_values(
    exercise_number: int, self_score: int, repo_url: str = _REPO_URL
) -> list[tuple[str, str]]:
    """Return ``(label-prefix, value)`` pairs for the top-level cover fields."""
    return [
        ("Submitting an exercise number", str(exercise_number)),
        ("Group ID code", "COSMOS77"),
        ("Recommendation for self-scoring", str(self_score)),
        ("Link to GITHUB", repo_url),
        ("A late submission confirmation", "no"),
    ]


def fill_docx(template: Path, output_docx: Path, *, self_score: int, exercise_number: int) -> Path:
    """Insert the field values into `template` and save to `output_docx`."""
    from docx import Document

    doc = Document(str(template))
    paras = list(doc.paragraphs)
    for para in paras:
        text = para.text.strip()
        for prefix, value in build_field_values(exercise_number, self_score):
            if text.startswith(prefix):
                _append_run(para, f" {value}")
                break
    for header, values in _STUDENTS:
        _fill_student(paras, header=header, values=values)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


def _append_run(paragraph, text: str) -> None:  # noqa: ANN001
    run = paragraph.add_run(text)
    if paragraph.runs and paragraph.runs[0].font.size:
        run.font.size = paragraph.runs[0].font.size


def _fill_student(paragraphs: Iterable, *, header: str, values: tuple[str, ...]) -> None:  # noqa: ANN001
    paras = list(paragraphs)
    for idx, para in enumerate(paras):
        if para.text.strip() == header:
            pairs = zip(_STUDENT_LABELS, values, strict=True)
            for offset, (label, value) in enumerate(pairs, start=1):
                target = paras[idx + offset]
                if target.text.strip().startswith(label):
                    _append_run(target, f" {value}")
            return


def convert_to_pdf(input_docx: Path, output_pdf: Path) -> Path:
    """Convert `input_docx` to PDF via docx2pdf, then a LibreOffice fallback."""
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    if _try_docx2pdf(input_docx, output_pdf) or _try_libreoffice(input_docx, output_pdf):
        return output_pdf
    raise RuntimeError(
        "no PDF tool available. Install Microsoft Word (docx2pdf) or "
        "`brew install --cask libreoffice` and re-run."
    )


def _try_docx2pdf(input_docx: Path, output_pdf: Path) -> bool:
    try:
        from docx2pdf import convert  # type: ignore[import-not-found]
    except ImportError:
        return False
    convert(str(input_docx), str(output_pdf))
    return output_pdf.exists()


def _try_libreoffice(input_docx: Path, output_pdf: Path) -> bool:
    tool = shutil.which("soffice") or shutil.which("libreoffice")
    if tool is None:
        return False
    cmd = [tool, "--headless", "--convert-to", "pdf", "--outdir"]
    cmd += [str(output_pdf.parent), str(input_docx)]
    try:
        subprocess.run(cmd, check=True, capture_output=True)
    except subprocess.CalledProcessError:
        return False
    produced = output_pdf.parent / (input_docx.stem + ".pdf")
    if produced.exists() and produced != output_pdf:
        produced.replace(output_pdf)
    return output_pdf.exists()


def verify_pdf(path: Path) -> None:
    """Validate the produced file is a real PDF (`%PDF` magic bytes)."""
    with path.open("rb") as fh:
        head = fh.read(5)
    if not head.startswith(b"%PDF-"):
        raise RuntimeError(f"{path} is not a valid PDF (got {head!r})")


def main(argv: list[str] | None = None) -> int:
    """Fill the cover template, convert to PDF, and verify the output."""
    parser = argparse.ArgumentParser(description="Generate the COSMOS77-ex06 cover PDF.")
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--self-score", type=int, default=85)
    parser.add_argument("--exercise-number", type=int, default=6)
    args = parser.parse_args(argv)
    docx_out = args.output.with_suffix(".filled.docx")
    fill_docx(
        args.template, docx_out, self_score=args.self_score, exercise_number=args.exercise_number
    )
    convert_to_pdf(docx_out, args.output)
    verify_pdf(args.output)
    print(f"OK: cover PDF written to {args.output.resolve()}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    sys.exit(main())
